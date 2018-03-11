from docx import Document
import mysql.connector
import os
from mysql.connector import errorcode
from DB import config
from DB.dictionary import Dictionary
from collections import Counter
from collections import deque


def order_query(query):
    precedence = {}
    precedence['!'] = 3
    precedence['&'] = 2
    precedence['|'] = 1
    precedence['('] = 0
    precedence[')'] = 0

    output = []
    operator_stack = []

    # while there are tokens to be read
    for token in query:

        # if left bracket
        if token == '(':
            operator_stack.append(token)

        # if right bracket, pop all operators from operator stack onto output until we hit left bracket
        elif token == ')':
            operator = operator_stack.pop()
            while operator != '(':
                output.append(operator)
                operator = operator_stack.pop()

        # if operator, pop operators from operator stack to queue if they are of higher precedence
        elif token in precedence:
            # if operator stack is not empty
            if operator_stack:
                current_operator = operator_stack[-1]
                while operator_stack and precedence[current_operator] > precedence[token]:
                    output.append(operator_stack.pop())
                    if operator_stack:
                        current_operator = operator_stack[-1]

            operator_stack.append(token)  # add token to stack

        # else if operands, add to output list
        else:
            output.append(token.lower())

    # while there are still operators on the stack, pop them into the queue
    while operator_stack:
        output.append(operator_stack.pop())
    return output


class DBHandler(object):
    def __init__(self):
        self._db = None
        self._cur = None
        self._dictionary = None

    def connect(self):
        try:
            self._db = mysql.connector.connect(**config.config)
            self._cur = self._db.cursor(buffered=True)
            self.__init_db__()
        except mysql.connector.Error as err:
            if err.errno == errorcode.ER_ACCESS_DENIED_ERROR:
                print "Bad username or password"
            elif err.errno == errorcode.ER_BAD_DB_ERROR:
                print "Database does not exist"
            else:
                print err
        else:
            print "db connected!"

    def __init_db__(self):
        for name, ddl in config.TABLES.iteritems():
            try:
                print "Creating table {}: ".format(name),
                self._cur.execute(ddl)
            except mysql.connector.Error as err:
                if err.errno == errorcode.ER_TABLE_EXISTS_ERROR:
                    print "already exists."
                else:
                    print(err.msg)
            else:
                print "OK"

    def build_dictionary(self):
        self._dictionary = Dictionary()
        self._cur.execute(
            "SELECT term, doc_id FROM emanueldb.posting_file"
        )
        posting_file_table = self._cur.fetchall()
        self._dictionary.build_dictionary_from_table(posting_file_table)

    def __insert_to_documents__(self, **kwargs):
        print "inserting to: Documents table"
        add_to = (
            "INSERT INTO documents "
            "(doc_id, title, author, brief, subject, location) "
            "VALUES (%s, %s, %s, %s, %s, %s)"
        )
        data = (kwargs['doc_id'], kwargs['title'], kwargs['author'], kwargs['brief'], kwargs['subject'], kwargs['location'])
        try:
            self._cur.execute(add_to, data)
            self._db.commit()
        except Exception as e:
            print "cannot insert to documents table: {}".format(e)

    def __insert_to_posting_file__(self, **kwargs):
        print "inserting to: posting file table"
        add_to = (
            "INSERT INTO posting_file "
            "(term, doc_id, hits) "
            "VALUES (%s, %s, %s)"
        )
        data = (kwargs['term'], kwargs['doc'], kwargs['hits'])
        try:
            self._cur.execute(add_to, data)
            self._db.commit()
        except Exception as e:
            print "cannot insert to posting_file table: {}".format(e)

    def __parse_file__(self, document, doc_id):
        words = document.paragraphs[1].text.lower().split()
        shows = Counter(words)
        words = list(set(words))
        for word in words:
            b = "!@#$'/." + '",'
            for char in b:
                word = word.replace(char, "")
            self.__insert_to_posting_file__(term=word, doc=doc_id, hits=shows[word])

    def load_new_doc(self, doc_path, title=None, author=None, subject=None):
        doc_id = self.__generate_id__()
        try:
            document = Document(doc_path)
            stored_files = "E:\Emanuel\storedFiles\{}.docx".format(doc_id)
            document.save(stored_files)
            text = document.paragraphs[1].text.split('\n')
            brief = ' '.join(text[:5])
            self.__insert_to_documents__(doc_id=doc_id, title=title, author=author, subject=subject, brief=brief, location=stored_files)
        except Exception as e:
            print "error loading Doc: {}".format(e)
            return
        else:
            self.__parse_file__(document, doc_id)
            os.remove(doc_path)
            print("File was transferred from inputFiles to stored_files")

    def __generate_id__(self):
        self._cur.execute("SELECT COUNT(*) FROM emanueldb.documents", )
        doc_id = self._cur.fetchall()[0][0]
        if doc_id:
            return int(doc_id) + 101
        else:
            return 101

    def __get_docs_by_id__(self, doc_list):
        docs = []
        for doc in doc_list:
            self._cur.execute(
                "SELECT * FROM emanueldb.documents WHERE doc_id={}".format(doc)
            )
            current_doc = self._cur.fetchall()[0]
            doc_id = current_doc[0]
            doc_title = current_doc[1]
            doc_author = current_doc[2]
            doc_subject = current_doc[3]
            doc_brief = current_doc[4]
            doc_location = current_doc[5]
            docs.append((doc_id, doc_title, doc_author, doc_subject, doc_brief, doc_location))
        return docs

    def __split_query_to_array(self, query):
        query = query.replace(")", "+)+")
        query = query.replace("(", "+(+")
        query = query.replace("!", "+!+")
        query = query.replace("&", "+&+")
        query = query.replace("|", "+|+")
        query = query.replace("%29", ")")
        query = query.replace("%7C", "|")
        query = query.replace("%26", "&")
        query = query.replace("%28", "(+")
        query = query.replace("%21", "!+ ")
        query = query.replace("%22", '+"+')
        query = query.replace("++", '+')
        query = query.split('+')
        query = filter(lambda char: char != '' and char != ' ', query)
        query_list = []
        for word in query:
            global word
            query_list.append(word.replace(" ", ""))
        words = filter(lambda a_word: a_word not in self._dictionary.operators, query)
        return query_list, words

    def __parse_query__(self, query):
        query_list, words = self.__split_query_to_array(query)
        query_queue = deque(order_query(query_list))
        final_result = []
        while query_queue:
            token = query_queue.popleft()
            token = token.lower()
            temp_result = []

            # if operand in dictionary
            if token not in self._dictionary.operators and token in self._dictionary.get_dictionary():
                temp_result.append(self._dictionary.find_in_dictionary(token))

            elif token == '&':  # if token == & (AND)
                try:
                    right_operand = final_result.pop()
                    left_operand = final_result.pop()
                    # print "{} & {}".format(right_operand, left_operand)   FOR DEBUG ONLY
                    temp_result = [self._dictionary.__execute_AND__(right_operand, left_operand)]
                except Exception as e:
                    print "error popping from queue: {}".format(e)
                    temp_result = [[]]

            elif token == '|':  # if token == | (OR)
                try:
                    right_operand = final_result.pop()
                    left_operand = final_result.pop()
                    # print "{} | {}".format(right_operand, left_operand)   FOR DEBUG ONLY
                    temp_result = [self._dictionary.__execute_OR__(right_operand, left_operand)]
                except Exception as e:
                    print "error popping from queue: {}".format(e)
                    temp_result = [[]]

            elif token == '!':  # if token == ! (NOT)
                try:
                    right_operand = final_result.pop()
                    # print "!{}".format(right_operand)  FOR DEBUG ONLY
                    temp_result = [self._dictionary.__execute_NOT__(right_operand)]
                except Exception as e:
                    print "error popping from queue: {}".format(e)
                    temp_result = [[]]
            else:
                temp_result.append([])
            final_result.append(temp_result[0])
        return final_result.pop(), words

    def search(self, query):
        result, words = self.__parse_query__(query)
        docs = set()
        if len(result) > 0:
            for doc in result:
                docs.add(doc)
        return self.__get_docs_by_id__(docs), words

    def close_connection(self):
        self._db.close()

# FOR DEBUG ONLY
# if __name__ == '__main__':
#     mydb = DBHandler()
#     mydb.connect()
#     mydb.build_dictionary()
#     print mydb.search("( need|     That ) & !forever")
